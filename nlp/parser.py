"""Извлечение текста из файлов резюме (PDF, DOCX, DOC, TXT)."""
import io
import logging

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    extractors = {
        "pdf": _from_pdf,
        "docx": _from_docx,
        "doc": _from_doc,
        "txt": _from_txt,
    }
    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Неподдерживаемый формат файла: .{ext}")
    text = extractor(file_bytes)
    logger.info("Extracted %d chars from %s", len(text), filename)
    return text


def _from_pdf(data: bytes) -> str:
    """
    Надёжное извлечение текста из PDF с поддержкой многоколонных макетов.
    Использует пословное извлечение с сортировкой по координатам.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        pages_text = []

        for page in doc:
            # Получаем все слова с координатами: (x0, y0, x1, y1, word, block, line, word_n)
            words = page.get_text("words")

            if not words:
                # Fallback: простое извлечение
                pages_text.append(page.get_text())
                continue

            # Группируем слова по строкам (бинуем y-координату с шагом 6 пикселей)
            lines: dict[int, list[tuple[float, str]]] = {}
            for w in words:
                y_bin = round(w[1] / 6)
                if y_bin not in lines:
                    lines[y_bin] = []
                lines[y_bin].append((w[0], w[4]))  # (x, слово)

            # Сортируем строки по y, слова внутри строки по x
            result_lines = []
            for y_key in sorted(lines.keys()):
                line_words = sorted(lines[y_key], key=lambda item: item[0])
                result_lines.append(" ".join(item[1] for item in line_words))

            pages_text.append("\n".join(result_lines))

        doc.close()
        result = "\n".join(pages_text).strip()

        if len(result) < 30:
            raise ValueError("PDF содержит слишком мало текста — возможно, файл отсканирован.")

        return result

    except ValueError:
        raise
    except Exception as e:
        logger.error("PDF parse error: %s", e)
        raise ValueError("Не удалось прочитать PDF-файл. Убедитесь, что файл не зашифрован.")


def _from_docx(data: bytes) -> str:
    """
    Извлечение текста из DOCX включая таблицы и текстовые блоки (textbox).
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(data))
        parts: list[str] = []

        # Обычные абзацы
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)

        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)

        # Текстовые блоки (textbox) — часто используются в шаблонах резюме
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            for para in txbx.iter(qn("w:p")):
                text = "".join(
                    node.text or ""
                    for node in para.iter(qn("w:t"))
                )
                if text.strip():
                    parts.append(text.strip())

        # Колонтитулы
        for section in doc.sections:
            for hdr in (section.header, section.footer):
                if hdr:
                    for p in hdr.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())

        result = "\n".join(parts).strip()
        if not result:
            raise ValueError("DOCX файл не содержит читаемого текста.")
        return result

    except ValueError:
        raise
    except Exception as e:
        logger.error("DOCX parse error: %s", e)
        raise ValueError("Не удалось прочитать DOCX-файл.")


def _from_doc(data: bytes) -> str:
    try:
        return _from_docx(data)
    except Exception:
        pass
    try:
        text = data.decode("cp1251", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if len(line.strip()) > 3]
        return "\n".join(lines[:500])
    except Exception as e:
        logger.error("DOC parse error: %s", e)
        raise ValueError(
            "Формат .doc имеет ограниченную поддержку. "
            "Конвертируйте файл в .docx или .pdf."
        )


def _from_txt(data: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Не удалось определить кодировку TXT-файла.")
